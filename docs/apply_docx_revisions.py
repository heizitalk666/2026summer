#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""按 D3 评审决议改《配电室设备状态视觉测控系统 设计方案书》。

    pip install python-docx
    python3 docs/apply_docx_revisions.py --check    # 只报告要改什么，不写
    python3 docs/apply_docx_revisions.py            # 写回 docx

**为什么用脚本而不是手工改 Word。** 决议一共动方案书 11 处文字加 2 个新增小节，
散在第 2、5、6、7、8、11 章。手工改的问题不是慢，是**改完没人能复核改了什么**——
docx 是二进制，git diff 只会显示"文件变了"。脚本本身就是修订说明的可执行版本，
逐条对应《一致性差异清单》第 7 节的决议编号，配套的人读版本在
`docs/方案书修订记录.md`。

脚本是**幂等**的：已经改过的条目会跳过并报告，重复运行不会改坏。
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
DOCX = REPO / "配电室设备状态视觉测控系统_设计方案书.docx"

# ---------------------------------------------------------------- 文字替换
#: (决议编号, 原文片段, 改后片段)。原文片段必须在文档中唯一命中。
REPLACEMENTS = [
    ("C8", "主控与云台共用一路 12 V 母线，经 DC-DC 降压至 5 V 供主控使用。",
     "主控与云台共用一路 24 V 母线，经两级 DC-DC 降压（24 V→12 V→5 V）供主控使用。"),

    ("C2", "状态机每 500 ms 向网关发送一次心跳，网关超过 1.5 s 未收到即判定上位失联，",
     "状态机每 200 ms 向网关发送一次心跳（5 Hz），网关超过 1.5 s 未收到即判定上位失联，"),

    ("C6", "通信协议中只定义四条任务级指令：暂停巡检、恢复巡检、缓慢前进、前往观测位。",
     "通信协议中只定义七条指令：暂停巡检、恢复巡检、缓慢前进、前往观测位、"
     "设定云台姿态、设定云台角速度、心跳。其中前四条是任务级运动指令，"
     "后三条不改变车辆运动。"),

    ("C9", "缓慢前进的距离不超过 0.5 m、速度不超过 0.2 m/s，观测位编号必须在预设列表内。",
     "缓慢前进的距离不超过 0.5 m，观测位编号必须在预设列表内。"
     "缓慢前进的速度上限由底盘侧常量保证，不由网关校验——协议里没有速度字段，"
     "网关拿不到速度就无从校验速度，这正是安全边界的体现。"),

    ("C3", "系统软件划分为六个独立进程，各自独立崩溃、独立重启，通过消息机制通信。",
     "系统软件划分为四个独立进程，各自独立崩溃、独立重启，通过消息机制通信。"
     "把需要访问同一帧的模块（检测、跟踪、质量评价、测量解算）合并在同一进程内，"
     "理由见本节末段的内存搬运分析；合并的代价是丢掉一个失效等级，"
     "由该进程内部的异常捕获补回轻失效路径（测量失败时降级归档，不触发看门狗）。"),

    ("B3", "单次测量产生的数据包含广角原图一张、三视角补拍图三张、"
     "触发前后各若干秒的视频片段与一份记录文件，合计约 6.7 MB。"
     "按单轮巡检 20 次测量计，单轮约 134 MB，每日六轮约 804 MB。"
     "64 GB 存储在完全不上传的极端情况下可缓存约 80 天，余量充足。",
     "单次测量产生的数据包含广角标注图与原图各一张、复核连拍图三张、"
     "读数 ROI 裁图一张与一份记录文件，合计约 1.7 MB；"
     "触发前后的视频片段首版不产出（接口上已预留 CRUISE_VIDEO 角色，见决议 B3），"
     "启用后单次约 6.7 MB。按单轮巡检 20 次测量计，不含视频单轮约 34 MB、"
     "每日六轮约 204 MB，64 GB 可缓存约 320 天；含视频则为 134 MB / 804 MB / 约 80 天。"
     "两种口径都留足余量，取值差额 5.0 MB 全部是视频。"),

    ("A3", "空间多视角。单次复核采集三个视角：主视角，以及左右各偏转 15° 的两个辅视角。"
     "三视角的作用有二，一是通过一致性比对抑制反光与偶发遮挡造成的误读，"
     "二是当主视角被高光覆盖时可用辅视角替代。",
     "空间多视角，条件式触发。单次复核默认在主视角连拍 3 帧（抗云台残余抖动造成的"
     "运动模糊）；当图像质量评价判定主视角存在高光遮挡，或首次读数置信度低于阈值时，"
     "才追加左右各偏转 15° 的两个辅视角。三视角的作用有二，一是通过一致性比对抑制"
     "反光与偶发遮挡造成的误读，二是当主视角被高光覆盖时可用辅视角替代。"
     "之所以做成条件式而不是每次都拍：辅视角要多花 1.5 s，而它真正有用的场合"
     "（镜面高光）只占少数，无条件付这个代价会把单轮复核次数从 21 次压到 18 次，"
     "用 14 % 的复核能力去换一个多数时候用不上的手段（决议 A3）。"),

    ("A3", "多视角一致性判定。三个视角的读数极差若超过 0.5 % FS，判定本次测量不可信，"
     "标记为质量受限并交人工复核，不写入台账。",
     "多视角一致性判定，仅在走了辅视角的那次复核上生效。三个视角的读数极差若超过 "
     "0.5 % FS，判定本次测量不可信，标记为质量受限并交人工复核，不写入台账；"
     "极差值随证据包上报（字段 multiview_spread）。只连拍不辅视角时该判定不适用，"
     "由质量评价的四项指标兜底。"),

    ("A3", "对三视角分别解算读数，做一致性判定与时域中值滤波，得到最终测量值",
     "对采集到的各视角分别解算读数，做时域中值滤波；若本次走了辅视角，"
     "再做三视角一致性判定，得到最终测量值"),
]

#: 表格里的整格替换：(决议编号, 单元格原文, 改后)
CELL_REPLACEMENTS = [
    # C4 + C10 的连锁：T_r 8.8 → 9.2 s。这个数在文档里出现两处（11.2 预期指标表、
    # 3.3.3 测量调整策略的时长核算），下面的替换是全文扫描，一次改干净。
    ("C4", "约 8.8 s", "约 9.2 s"),
    ("C3", "六进程全开连续运行 30 分钟", "四进程全开连续运行 30 分钟"),
    ("C3", "六进程部署", "四进程部署"),
    ("C3", "RK3576 平台，六进程系统已部署并配置开机自启",
     "RK3576 平台，四进程系统已部署并配置开机自启"),
    ("C3", "在此基础上加入指示灯与开关位置识别，六进程在板端实测通过，安全机制三项演示通过",
     "在此基础上加入指示灯与开关位置识别，四进程在板端实测通过，安全机制三项演示通过"),
    ("C7", "四组接口的数据结构与约定", "四条接口共五份 JSON Schema 的数据结构与约定"),
]

#: 表 6-5 状态超时：以 ICD 为准（决议 C1）。方案书原表九个状态里有八个与 ICD 不同，
#: 其中 CAPTURE 还要因 A3 的条件路径从 1.5 s 放宽到 4.0 s。
TABLE_6_5_TIMEOUTS = {
    "SUSPECT": "0.5 s", "HALT_REQ": "4.0 s", "AIM": "3.0 s", "ZOOM": "2.5 s",
    "CAPTURE": "4.0 s", "VERIFY": "5.0 s", "PACK": "2.0 s", "RESUME": "1.0 s",
}

#: 表 8-1 进程划分：六行改四行（C3）
TABLE_8_1_ROWS = [
    ("感知", "图像采集、目标检测、跟踪、质量评价、测量解算（标度变换、滤波、透视校正）",
     "NPU + A72",
     "看门狗触发，车辆恢复巡航；测量解算单独失败时由进程内异常捕获降级归档，不触发看门狗"),
    ("任务", "复核流程决策、云台 PID、复核预算调度、超时管理", "A53", "看门狗触发"),
    ("网关", "指令白名单校验、参数范围硬校验、心跳监测、审计日志、底盘协议转换", "A53",
     "网关重启期间拒绝一切运动指令；车辆按底盘安全层自行处置"),
    ("归档", "证据打包、校验、本地队列、断点续传上传", "A53", "数据暂存本地，不影响测控"),
]


def _replace_in_paragraph(p, old: str, new: str) -> bool:
    """把段落里的 old 换成 new，尽量保住格式。

    docx 会把一段话切成多个 run（改过字体、加过粗的地方就断一次），要找的
    字符串常常横跨几个 run。这里的做法是：命中后把整段文字写进第一个 run、
    其余 run 清空。段落级的格式（样式、缩进、编号）不受影响，段内的局部
    格式差异会被抹平——正文段落本来就是一致的字体，可以接受。
    """
    if old not in p.text:
        return False
    text = p.text.replace(old, new)
    if not p.runs:
        return False
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""
    return True


def _iter_paragraphs(doc):
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _insert_after(paragraph, text: str, style: str):
    """在 paragraph 之后插一段，返回新段落。"""
    import copy
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, paragraph._parent)
    for r in para.runs[1:]:
        r._r.getparent().remove(r._r)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)
    para.style = paragraph.part.document.styles[style]
    return para


# ---------------------------------------------------------------- 新增小节
NEW_SECTIONS = [
    # (决议, 插在哪个标题所属小节之后, [(样式, 文字), ...])
    ("B1", "6.2.2　检测模型选型", [
        ("Heading 3", "6.2.3　未知异常检测"),
        ("Normal",
         "前两小节处理的是「已知缺陷」——类别事先定义、有标注数据、可以监督训练。"
         "但配电室里真正要防的一类问题恰恰不是这样：渗漏油、呼吸器变色、异物侵入、"
         "接线端子发黑，这些既没有公开标注数据，室内场景的样本更是几乎为零。"
         "2.4.5 节把这一条列为数据来源约束，本节给出绕开它的办法。"),
        ("Normal",
         "采用非监督异常检测：只用「看起来正常」的样本训练，模型学习正常样本的特征分布，"
         "推理时把偏离该分布的区域标为异常，不需要任何缺陷标注。"
         "正常样本的可得性与缺陷样本相反——巡检跑一轮就攒下一批，"
         "而且它们来自本场景本设备，分布比任何公开数据集都贴合。"),
        ("Normal",
         "本方案实现两条通路并做比选。基线是统计法：在线学习正常样本的特征均值与协方差，"
         "按马氏距离打异常分，零权重、当场可跑，且异常分来自哪个特征通道说得清——"
         "可解释性在验收场合是实打实的加分项。对照组是 PaDiM / EfficientAD 一类"
         "基于预训练主干的方法，精度更高但需要权重、不可解释。"
         "哪一条进最终系统由同一批样本上的误报率与漏报率决定，"
         "比选结论是「不启用学习法」也是成果。"),
        ("Normal",
         "这一路在系统里的位置是 L3：L1 负责「看见有东西」，L2 负责「读出数值」，"
         "L3 负责「说不出是什么但不对劲」。它同时承接了首版缺陷类别收窄"
         "（决议 A2 把渗漏油换成开关分合位）之后空出来的那条纯检测通路——"
         "渗漏油不再作为一个监督类别去训，而是作为「未知异常」被 L3 兜住。"
         "L3 的输出不单独定案，进入 L4 显式仲裁与其余三路证据一起决策。"),
    ]),
    ("B2", "6.3.6　车辆运动的逻辑控制", [
        ("Heading 3", "6.3.7　复核预算与失锁抑制"),
        ("Normal",
         "主动复核会停车，停车就吃掉巡检时间。单轮巡检的时间上限 T_max 是硬的，"
         "巡航本身要占 L/v，剩下的才归复核，所以一轮里能做多少次复核是有上限的："
         "N_max = ⌊(T_max − L/v) / T_r⌋。按 L = 200 m、v = 0.5 m/s、T_max = 600 s、"
         "单次复核 T_r = 9.2 s 计，N_max = 21 次。"
         "不设这个预算的后果不是慢，是巡检跑不完——车会停在半路上反复复核，"
         "而路线尽头的设备一次也没看过。"),
        ("Normal",
         "预算耗尽后不丢弃触发：可疑标记照常置位，事件进入顺延队列按优先级排序，"
         "下一轮巡检优先处理。优先级 = 严重度 × 置信度 × 新颖度，"
         "其中新颖度对复现的缺陷取 0.3 而不是 0——取 0 会导致某个缺陷第一轮复核失败之后"
         "永远排不上队。"),
        ("Normal",
         "另有三条抑制规则挡住三类不同的死循环：同一跟踪 ID 复核过后 60 s 内不再复核，"
         "挡的是同一个目标反复触发；同一巡检位 2 m 半径内本轮只复核一次，"
         "挡的是跟踪器丢了 ID 之后同一目标以新 ID 再次触发；恢复巡航后 3 s 静默，"
         "挡的是车刚起步、云台刚归位那一瞬间画面剧烈变化引发的连锁触发。"
         "少任何一条都有一类循环补不上。"),
        ("Normal",
         "定位失锁时一律不进复核。复核的前提是「知道自己在哪」——巡检位半径抑制、"
         "证据包里的位姿、顺延队列的排序全都依赖定位，位姿无效时这三样同时失去意义，"
         "此时继续复核只会产出无法归档的证据。失锁按抑制处理而不是按故障处理："
         "车照常巡航，只是不做主动测量，定位恢复后自动恢复复核能力。"),
    ]),
]


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="按 D3 决议改设计方案书")
    ap.add_argument("--check", action="store_true", help="只报告，不写回")
    a = ap.parse_args(argv)

    try:
        from docx import Document
    except ImportError:
        print("需要 python-docx：pip install python-docx", file=sys.stderr)
        return 2

    doc = Document(str(DOCX))
    done, skipped = [], []

    for tag, old, new in REPLACEMENTS:
        hit = False
        for p in _iter_paragraphs(doc):
            if _replace_in_paragraph(p, old, new):
                hit = True
                break
        (done if hit else skipped).append((tag, old[:36]))

    for tag, old, new in CELL_REPLACEMENTS:
        # 这一类是短片段，可能在多处出现（如"约 8.8 s"），全文扫完不要提前退出，
        # 否则要跑好几遍才改干净，而"跑几遍结果不一样"的脚本没法用来复核。
        n = sum(1 for p in _iter_paragraphs(doc) if _replace_in_paragraph(p, old, new))
        (done if n else skipped).append((tag, "%s（%d 处）" % (old[:30], n)))

    # 表 6-5：状态超时以 ICD 为准（C1）
    for t in doc.tables:
        head = [c.text.strip() for c in t.rows[0].cells]
        if head[:2] == ["状态", "含义"] and "超时" in head:
            col = head.index("超时")
            n = 0
            for row in t.rows[1:]:
                name = row.cells[0].text.strip()
                want = TABLE_6_5_TIMEOUTS.get(name)
                cell = row.cells[col]
                if want and cell.text.strip() != want:
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = want
                        for r in cell.paragraphs[0].runs[1:]:
                            r.text = ""
                    else:
                        cell.text = want
                    n += 1
            # CAPTURE 那一行的含义也要跟 A3 对上
            for row in t.rows[1:]:
                if row.cells[0].text.strip() == "CAPTURE":
                    _replace_in_paragraph(row.cells[1].paragraphs[0],
                                          "三视角采集", "主视角连拍，必要时追加辅视角")
            (done if n else skipped).append(("C1", "表 6-5 状态超时（%d 格）" % n))
            break
    else:
        skipped.append(("C1", "表 6-5"))

    # 表 8-1：六行改四行
    for t in doc.tables:
        head = [c.text.strip() for c in t.rows[0].cells]
        if head[:2] == ["进程", "职责"] and len(t.rows) == 7:
            for row, vals in zip(t.rows[1:], TABLE_8_1_ROWS):
                for cell, v in zip(row.cells, vals):
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = v
                        for r in cell.paragraphs[0].runs[1:]:
                            r.text = ""
                    else:
                        cell.text = v
            for row in list(t.rows)[5:]:
                row._tr.getparent().remove(row._tr)
            done.append(("C3", "表 8-1 六进程 → 四进程"))
            break
    else:
        skipped.append(("C3", "表 8-1"))

    # 新增小节
    for tag, after_heading, blocks in NEW_SECTIONS:
        if any(blocks[0][1] in p.text for p in doc.paragraphs):
            skipped.append((tag, blocks[0][1] + "（已存在）"))
            continue
        idx = next((i for i, p in enumerate(doc.paragraphs)
                    if p.text.strip().startswith(after_heading.split("　")[0])), None)
        if idx is None:
            skipped.append((tag, after_heading))
            continue
        # 插到该小节最后一段之后（下一个标题之前）
        j = idx + 1
        while j < len(doc.paragraphs) and not doc.paragraphs[j].style.name.startswith("Heading"):
            j += 1
        anchor = doc.paragraphs[j - 1]
        for style, text in blocks:
            anchor = _insert_after(anchor, text, style)
        done.append((tag, blocks[0][1]))

    print("已改 %d 处：" % len(done))
    for tag, what in done:
        print("  [%s] %s" % (tag, what))
    if skipped:
        print("\n跳过 %d 处（已改过或没命中）：" % len(skipped))
        for tag, what in skipped:
            print("  [%s] %s" % (tag, what))

    if a.check:
        print("\n--check：未写回")
        return 0
    if done:
        doc.save(str(DOCX))
        print("\n已写回 %s" % DOCX.name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
